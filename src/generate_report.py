"""
Generador del informe en Word con los resultados del modelo.

Reúne las métricas del entrenamiento, las gráficas de evaluación y, sobre
todo, las imágenes de las predicciones realizadas sobre las imágenes propias
(imagen procesada + etiqueta real + etiqueta predicha + probabilidad).

Requisitos previos:
    python src/train_model.py        -> modelo/metricas.json
    python src/predict_images.py     -> resultados/predicciones.json

Uso:
    python src/generate_report.py
"""

from __future__ import annotations

import json
from datetime import datetime

import matplotlib

matplotlib.use("Agg")
import matplotlib.pyplot as plt
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

from config import (
    DIR_IMAGENES_PROPIAS,
    DIR_RESULTADOS,
    NOMBRE_DATASET,
    NUM_CLASES,
    RUTA_METRICAS,
    RUTA_MODELO,
)

RUTA_INFORME = DIR_RESULTADOS / "Informe_Reconocimiento_Letras_EMNIST.docx"
RUTA_PREDICCIONES = DIR_RESULTADOS / "predicciones.json"


# --------------------------------------------------------------------------
# Utilidades de documento
# --------------------------------------------------------------------------
def pie_de_figura(doc, texto):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(texto)
    run.italic = True
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x55, 0x5F, 0x6D)
    return p


def insertar_imagen(doc, ruta, ancho=6.2, pie=None):
    if not ruta or not ruta.exists():
        doc.add_paragraph(f"[No se encontró la imagen: {ruta}]")
        return
    doc.add_picture(str(ruta), width=Inches(ancho))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    if pie:
        pie_de_figura(doc, pie)


def tabla_clave_valor(doc, filas, encabezados=("Concepto", "Valor")):
    tabla = doc.add_table(rows=1, cols=2)
    tabla.style = "Light Grid Accent 1"
    for celda, texto in zip(tabla.rows[0].cells, encabezados):
        celda.text = texto
    for clave, valor in filas:
        celda_a, celda_b = tabla.add_row().cells
        celda_a.text = str(clave)
        celda_b.text = str(valor)
    return tabla


# --------------------------------------------------------------------------
# Figura auxiliar: los pasos del preprocesamiento
# --------------------------------------------------------------------------
def figura_pasos_preprocesamiento(ruta_salida):
    """Ilustra los cuatro pasos del preprocesamiento sobre una imagen propia."""
    from preprocess import preprocesar_imagen

    candidatas = sorted(DIR_IMAGENES_PROPIAS.glob("*.png")) or sorted(
        DIR_IMAGENES_PROPIAS.glob("*.jpg")
    )
    if not candidatas:
        return None
    ejemplo = candidatas[0]
    _, pasos = preprocesar_imagen(ejemplo, devolver_pasos=True)

    fig, ejes = plt.subplots(1, 4, figsize=(12, 3.2))
    ejes[0].imshow(pasos["original"])
    ejes[0].set_title("1. Original\n(color, tamaño libre)", fontsize=10)

    ejes[1].imshow(pasos["gris"], cmap="gray", vmin=0, vmax=255)
    ejes[1].set_title("2. Escala de grises", fontsize=10)

    ejes[2].imshow(pasos["invertida"], cmap="gray", vmin=0, vmax=255)
    invertida = "aplicada" if pasos["se_invirtio"] else "no necesaria"
    ejes[2].set_title(f"3. Inversión de colores\n({invertida})", fontsize=10)

    ejes[3].imshow(pasos["normalizada"], cmap="gray", vmin=0, vmax=1)
    ejes[3].set_title("4. Redimensionado 28×28\n+ normalización [0,1]", fontsize=10)

    for ax in ejes:
        ax.axis("off")
    fig.suptitle(f"Preprocesamiento aplicado a «{ejemplo.name}»", fontsize=12)
    fig.tight_layout()
    fig.savefig(ruta_salida, dpi=140, bbox_inches="tight")
    plt.close(fig)
    return ruta_salida


# --------------------------------------------------------------------------
# Secciones del informe
# --------------------------------------------------------------------------
def portada(doc, metricas):
    for _ in range(4):
        doc.add_paragraph()
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = t.add_run("Reconocimiento de letras manuscritas\ncon redes neuronales y EMNIST")
    run.bold = True
    run.font.size = Pt(26)

    s = doc.add_paragraph()
    s.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = s.add_run("Examen final práctico — Informe de resultados")
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x55, 0x5F, 0x6D)

    doc.add_paragraph()
    d = doc.add_paragraph()
    d.alignment = WD_ALIGN_PARAGRAPH.CENTER
    exactitud = metricas.get("exactitud_test")
    detalle = f"Exactitud sobre el conjunto de prueba EMNIST: {exactitud:.2%}" if exactitud else ""
    run = d.add_run(
        f"Modelo: {metricas.get('arquitectura', '—').upper()}   |   "
        f"Dataset: {metricas.get('dataset', NOMBRE_DATASET)}\n{detalle}\n\n"
        f"{datetime.now().strftime('%d/%m/%Y')}"
    )
    run.font.size = Pt(11)
    doc.add_page_break()


def seccion_objetivo(doc):
    doc.add_heading("1. Objetivo y alcance", level=1)
    doc.add_paragraph(
        "El trabajo consiste en entrenar una red neuronal capaz de reconocer letras "
        "manuscritas a partir del conjunto de datos EMNIST y en aplicarla después "
        "sobre imágenes propias, ajenas al dataset. El sistema se completa con una "
        "página web que ofrece un recuadro donde escribir una letra y un servidor "
        "en Python que responde con la predicción del modelo ya entrenado."
    )
    doc.add_paragraph("Componentes entregados:", style="Intense Quote")
    for texto in [
        "Entrenamiento y evaluación de la red sobre EMNIST (src/train_model.py).",
        "Preprocesamiento de imágenes propias (src/preprocess.py).",
        "Predicción sobre las imágenes propias (src/predict_images.py).",
        "Servidor de predicción en Python con Flask (src/server.py).",
        "Página web con el recuadro de escritura y el resultado (web/index.html).",
        "Este informe, generado automáticamente (src/generate_report.py).",
    ]:
        doc.add_paragraph(texto, style="List Bullet")


def seccion_dataset(doc, metricas):
    doc.add_heading("2. Conjunto de datos EMNIST", level=1)
    doc.add_paragraph(
        "Se emplea la variante EMNIST «letters», documentada en el catálogo de "
        "TensorFlow Datasets (https://www.tensorflow.org/datasets/catalog/emnist). "
        "Contiene caracteres manuscritos repartidos en 26 clases balanceadas, una "
        "por cada letra del alfabeto; las mayúsculas y minúsculas de una misma "
        "letra comparten etiqueta."
    )
    tabla_clave_valor(
        doc,
        [
            ("Dataset", metricas.get("dataset", NOMBRE_DATASET)),
            ("Imágenes de entrenamiento", f"{metricas.get('imagenes_entrenamiento', 0):,}"),
            ("Imágenes de prueba", f"{metricas.get('imagenes_prueba', 0):,}"),
            ("Número de clases", NUM_CLASES),
            ("Formato de imagen", "28 × 28 píxeles, escala de grises"),
            ("Convención de color", "Trazo blanco sobre fondo negro"),
        ],
    )
    doc.add_paragraph()
    doc.add_paragraph(
        "Nota técnica: las imágenes de EMNIST se distribuyen traspuestas (giradas y "
        "espejadas) respecto a su orientación natural. En la carga del dataset se "
        "aplica una trasposición para dejarlas derechas, de modo que coincidan con "
        "la orientación de las imágenes propias que se procesan más adelante."
    )


def seccion_arquitectura(doc, metricas):
    doc.add_heading("3. Arquitectura de la red neuronal", level=1)
    doc.add_paragraph(
        "La red cumple la estructura solicitada: una capa de entrada, varias capas "
        "ocultas (más de las dos exigidas) y una capa de salida con activación "
        "softmax de 26 unidades, una por letra."
    )

    try:
        import tensorflow as tf

        modelo = tf.keras.models.load_model(RUTA_MODELO)
        tabla = doc.add_table(rows=1, cols=4)
        tabla.style = "Light Grid Accent 1"
        for celda, texto in zip(
            tabla.rows[0].cells, ("Capa", "Tipo", "Forma de salida", "Parámetros")
        ):
            celda.text = texto
        for capa in modelo.layers:
            fila = tabla.add_row().cells
            fila[0].text = capa.name
            fila[1].text = type(capa).__name__
            try:
                fila[2].text = str(tuple(capa.output.shape))
            except Exception:
                fila[2].text = "—"
            fila[3].text = f"{capa.count_params():,}"
        doc.add_paragraph()
        doc.add_paragraph(f"Parámetros entrenables totales: {modelo.count_params():,}")
    except Exception as exc:
        doc.add_paragraph(f"[No se pudo cargar el modelo para detallar las capas: {exc}]")

    doc.add_paragraph()
    tabla_clave_valor(
        doc,
        [
            ("Optimizador", "Adam (learning rate inicial 1e-3)"),
            ("Función de pérdida", "sparse_categorical_crossentropy"),
            ("Métrica", "accuracy"),
            ("Tamaño de lote", 128),
            ("Épocas ejecutadas", metricas.get("epocas_ejecutadas", "—")),
            ("Regularización", "BatchNormalization + Dropout"),
            ("Aumento de datos", "Rotación ±12°, traslación ±2 px, zoom ±10 %"),
        ],
        encabezados=("Hiperparámetro", "Valor"),
    )


def seccion_entrenamiento(doc, metricas):
    doc.add_heading("4. Entrenamiento y evaluación sobre EMNIST", level=1)
    arq = metricas.get("arquitectura", "cnn")
    doc.add_paragraph(
        "El modelo se entrenó reservando un 10 % del split de entrenamiento para "
        "validación, con parada temprana sobre la exactitud de validación y "
        "reducción automática del learning rate. La evaluación final se hizo sobre "
        "el conjunto de prueba de EMNIST, que el modelo no vio en ningún momento."
    )

    tabla_clave_valor(
        doc,
        [
            ("Exactitud en prueba", f"{metricas.get('exactitud_test', 0):.2%}"),
            ("Exactitud top-3 en prueba", f"{metricas.get('exactitud_top3_test', 0):.2%}"),
            ("Pérdida en prueba", f"{metricas.get('perdida_test', 0):.4f}"),
            ("Exactitud en validación", f"{metricas.get('exactitud_val_final', 0):.2%}"),
            ("Épocas ejecutadas", metricas.get("epocas_ejecutadas", "—")),
            ("Duración del entrenamiento", f"{metricas.get('duracion_minutos', 0)} minutos"),
        ],
        encabezados=("Métrica", "Resultado"),
    )

    doc.add_paragraph()
    insertar_imagen(
        doc,
        DIR_RESULTADOS / f"historial_entrenamiento_{arq}.png",
        pie="Figura 1. Evolución de la exactitud y la pérdida durante el entrenamiento.",
    )
    doc.add_page_break()
    insertar_imagen(
        doc,
        DIR_RESULTADOS / f"matriz_confusion_{arq}.png",
        ancho=5.6,
        pie="Figura 2. Matriz de confusión normalizada sobre el conjunto de prueba de EMNIST.",
    )

    por_clase = metricas.get("exactitud_por_clase", {})
    if por_clase:
        peores = sorted(por_clase.items(), key=lambda kv: kv[1])[:5]
        doc.add_paragraph()
        doc.add_paragraph(
            "Las clases con más confusión son las que comparten trazo en manuscrita: "
            + ", ".join(f"{c} ({v:.1%})" for c, v in peores)
            + "."
        )

    doc.add_page_break()
    insertar_imagen(
        doc,
        DIR_RESULTADOS / f"muestras_test_{arq}.png",
        pie="Figura 3. Ejemplos de predicción sobre el conjunto de prueba de EMNIST "
        "(verde = acierto, rojo = fallo).",
    )


def seccion_preprocesamiento(doc):
    doc.add_heading("5. Preprocesamiento de las imágenes propias", level=1)
    doc.add_paragraph(
        "Las imágenes propias no tienen el formato de EMNIST: llegan en color, con "
        "tamaños distintos y normalmente con tinta oscura sobre papel claro. El "
        "módulo de preprocesamiento las adapta aplicando los pasos siguientes."
    )
    for texto in [
        "Conversión a escala de grises (un único canal de intensidad).",
        "Inversión de colores condicional: se mide el brillo del borde de la imagen "
        "y, si el fondo es claro, se invierte para dejar trazo blanco sobre fondo "
        "negro, que es el convenio de EMNIST.",
        "Redimensionamiento a 28 × 28 píxeles. El carácter se recorta a su caja "
        "mínima, se escala conservando la proporción hasta caber en una caja de "
        "20 × 20 y se centra por su centro de masa, replicando el criterio con el "
        "que se construyó el propio EMNIST.",
        "Normalización de los valores de los píxeles al rango [0, 1].",
    ]:
        doc.add_paragraph(texto, style="List Number")

    ruta_pasos = DIR_RESULTADOS / "pasos_preprocesamiento.png"
    figura_pasos_preprocesamiento(ruta_pasos)
    doc.add_paragraph()
    insertar_imagen(
        doc, ruta_pasos, pie="Figura 4. Los pasos del preprocesamiento sobre una imagen propia."
    )


def seccion_predicciones(doc, datos):
    doc.add_heading("6. Predicción sobre las imágenes propias", level=1)
    predicciones = datos.get("predicciones", [])
    exactitud = datos.get("exactitud_imagenes_propias")

    doc.add_paragraph(
        f"Se procesaron {datos.get('total_imagenes', 0)} imágenes propias, ajenas al "
        "conjunto EMNIST. Para cada una se muestra la imagen ya procesada, la "
        "etiqueta real, la etiqueta predicha por el modelo y la probabilidad "
        "asociada a esa predicción."
    )
    if exactitud is not None:
        doc.add_paragraph(
            f"Resultado global: {datos.get('aciertos', 0)} aciertos de "
            f"{datos.get('total_con_etiqueta', 0)} imágenes etiquetadas "
            f"({exactitud:.2%} de acierto).",
            style="Intense Quote",
        )

    insertar_imagen(
        doc,
        DIR_RESULTADOS / "predicciones_propias.png",
        pie="Figura 5. Resumen de las predicciones sobre las imágenes propias "
        "(borde verde = acierto, borde rojo = fallo).",
    )

    doc.add_page_break()
    doc.add_heading("6.1. Tabla de resultados", level=2)
    tabla = doc.add_table(rows=1, cols=5)
    tabla.style = "Light Grid Accent 1"
    for celda, texto in zip(
        tabla.rows[0].cells,
        ("Archivo", "Etiqueta real", "Etiqueta predicha", "Probabilidad", "Resultado"),
    ):
        celda.text = texto
    for p in predicciones:
        fila = tabla.add_row().cells
        fila[0].text = p["archivo"]
        fila[1].text = p["etiqueta_real"] or "—"
        fila[2].text = p["etiqueta_predicha"]
        fila[3].text = f"{p['probabilidad']:.2%}"
        if p["etiqueta_real"] is None:
            fila[4].text = "—"
        else:
            fila[4].text = "Correcto" if p["etiqueta_predicha"] == p["etiqueta_real"] else "Incorrecto"

    doc.add_page_break()
    doc.add_heading("6.2. Detalle imagen por imagen", level=2)
    doc.add_paragraph(
        "Cada ficha muestra la imagen original, la imagen procesada que recibe la "
        "red y las cinco clases más probables con su probabilidad."
    )
    from pathlib import Path

    for i, p in enumerate(predicciones, start=1):
        estado = (
            "—"
            if p["etiqueta_real"] is None
            else ("acierto" if p["etiqueta_predicha"] == p["etiqueta_real"] else "fallo")
        )
        doc.add_heading(
            f"{p['archivo']} — real: {p['etiqueta_real'] or '?'} | "
            f"predicha: {p['etiqueta_predicha']} | {p['probabilidad']:.2%} ({estado})",
            level=3,
        )
        insertar_imagen(doc, Path(p["detalle"]), ancho=6.0)
        if i % 2 == 0:
            doc.add_page_break()


def seccion_aplicacion(doc):
    doc.add_heading("7. Página web y servidor de predicción", level=1)
    doc.add_paragraph(
        "La página «web/index.html» presenta un recuadro donde se escribe una letra "
        "con el ratón o con el dedo, y justo debajo muestra el resultado: la letra "
        "predicha, su probabilidad, la lista de las cinco candidatas más probables "
        "y la miniatura 28 × 28 que realmente recibe la red."
    )
    doc.add_paragraph(
        "El servidor «src/server.py» (Flask) carga el modelo entrenado una sola vez "
        "y expone el endpoint POST /predecir. La imagen recibida se procesa con el "
        "mismo módulo de preprocesamiento usado con las imágenes propias, lo que "
        "garantiza que el formato de entrada sea idéntico al del entrenamiento."
    )
    doc.add_paragraph("Puesta en marcha:", style="Intense Quote")
    for texto in [
        "python src/server.py",
        "Abrir http://127.0.0.1:5000 en el navegador.",
    ]:
        doc.add_paragraph(texto, style="List Bullet")

    captura = DIR_RESULTADOS / "captura_web.png"
    if captura.exists():
        doc.add_paragraph()
        insertar_imagen(
            doc,
            captura,
            ancho=5.2,
            pie="Figura 6. Página web prediciendo una letra escrita en el recuadro.",
        )


def seccion_conclusiones(doc, metricas, datos):
    doc.add_heading("8. Conclusiones", level=1)
    exactitud_emnist = metricas.get("exactitud_test", 0)
    exactitud_propias = datos.get("exactitud_imagenes_propias")

    doc.add_paragraph(
        f"La red alcanza un {exactitud_emnist:.2%} de exactitud sobre el conjunto de "
        f"prueba de EMNIST, un resultado coherente con el estado del arte para las "
        f"26 clases de esta variante, donde parte del error es irreducible: letras "
        f"como I/L, G/Q o U/V resultan ambiguas incluso para un lector humano cuando "
        f"se escriben aisladas y sin contexto."
    )
    if exactitud_propias is not None:
        doc.add_paragraph(
            f"Sobre las imágenes propias, ajenas al dataset, la exactitud es del "
            f"{exactitud_propias:.2%}. La diferencia respecto a EMNIST se explica por "
            f"el cambio de dominio: EMNIST recoge escritura manuscrita real, mientras "
            f"que las imágenes propias empleadas aquí son caracteres renderizados y "
            f"fotografiados sintéticamente."
        )
    doc.add_paragraph(
        "El paso decisivo para que el modelo funcione fuera de EMNIST no está en la "
        "arquitectura sino en el preprocesamiento: sin la inversión condicional de "
        "color y sin el centrado por centro de masa dentro de una caja de 20 × 20, "
        "las imágenes propias quedan fuera de la distribución con la que se entrenó "
        "la red y la exactitud cae drásticamente."
    )


# --------------------------------------------------------------------------
def main():
    if not RUTA_METRICAS.exists():
        raise SystemExit(
            f"Faltan las métricas ({RUTA_METRICAS}). Ejecuta antes: python src/train_model.py"
        )
    if not RUTA_PREDICCIONES.exists():
        raise SystemExit(
            f"Faltan las predicciones ({RUTA_PREDICCIONES}). "
            "Ejecuta antes: python src/predict_images.py"
        )

    metricas = json.loads(RUTA_METRICAS.read_text(encoding="utf-8"))
    datos = json.loads(RUTA_PREDICCIONES.read_text(encoding="utf-8"))

    doc = Document()
    seccion = doc.sections[0]
    seccion.orientation = WD_ORIENT.PORTRAIT
    for margen in ("top_margin", "bottom_margin", "left_margin", "right_margin"):
        setattr(seccion, margen, Inches(0.85))

    estilo = doc.styles["Normal"]
    estilo.font.name = "Calibri"
    estilo.font.size = Pt(11)

    portada(doc, metricas)
    seccion_objetivo(doc)
    seccion_dataset(doc, metricas)
    doc.add_page_break()
    seccion_arquitectura(doc, metricas)
    doc.add_page_break()
    seccion_entrenamiento(doc, metricas)
    doc.add_page_break()
    seccion_preprocesamiento(doc)
    doc.add_page_break()
    seccion_predicciones(doc, datos)
    doc.add_page_break()
    seccion_aplicacion(doc)
    seccion_conclusiones(doc, metricas, datos)

    doc.save(RUTA_INFORME)
    print(f"Informe generado: {RUTA_INFORME}")


if __name__ == "__main__":
    main()
